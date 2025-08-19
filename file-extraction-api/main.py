from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import uvicorn
import tempfile
import os
from pathlib import Path
import logging

# File processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = True

try:
    import json
    JSON_AVAILABLE = True
except ImportError:
    JSON_AVAILABLE = True

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="File Text Extraction API",
    description="Extract text from files for PII detection",
    version="1.0.0"
)

# Enable CORS for Chrome extension
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF files with multiple fallback methods"""
    if not PDF_AVAILABLE:
        return "[PDF file - PyPDF2 not installed. Install with: pip install PyPDF2]"
    
    extracted_text = ""
    errors = []
    
    # Method 1: Try PyPDF2 with different configurations
    try:
        logger.info("Trying PyPDF2 extraction...")
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            logger.info(f"PDF has {len(pdf_reader.pages)} pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        extracted_text += f"Page {page_num + 1}:\n{page_text}\n\n"
                        logger.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                    else:
                        logger.warning(f"Page {page_num + 1} appears to be empty or unreadable")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {e}")
                    errors.append(f"Page {page_num + 1}: {e}")
            
            if extracted_text.strip():
                logger.info(f"Successfully extracted {len(extracted_text)} characters total")
                return extracted_text.strip()
            else:
                logger.warning("No text could be extracted from any pages")
                
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {e}")
        errors.append(f"PyPDF2: {e}")
    
    # Method 2: Try alternative approach with strict=False
    try:
        logger.info("Trying PyPDF2 with strict=False...")
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file, strict=False)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        extracted_text += f"Page {page_num + 1}:\n{page_text}\n\n"
                except Exception as e:
                    logger.warning(f"Page {page_num + 1} error: {e}")
            
            if extracted_text.strip():
                return extracted_text.strip()
                
    except Exception as e:
        logger.error(f"PyPDF2 strict=False failed: {e}")
        errors.append(f"PyPDF2 (strict=False): {e}")
    
    # Method 3: Try to extract raw text using basic text extraction
    try:
        logger.info("Trying basic text extraction...")
        with open(file_path, 'rb') as file:
            content = file.read()
            
            # Look for text patterns in the binary data
            text_parts = []
            content_str = content.decode('latin-1', errors='ignore')
            
            # Extract readable text sequences
            import re
            readable_text = re.findall(r'[a-zA-Z0-9\s\.\,\;\:\!\?\-\@\#\$\%\^\&\*\(\)]{10,}', content_str)
            
            if readable_text:
                extracted_text = '\n'.join(readable_text)
                logger.info(f"Basic extraction found {len(extracted_text)} characters")
                return f"[Basic extraction from PDF]\n{extracted_text}"
                
    except Exception as e:
        logger.error(f"Basic text extraction failed: {e}")
        errors.append(f"Basic extraction: {e}")
    
    # If all methods fail, return detailed error
    error_msg = f"[PDF extraction failed - tried multiple methods]\nErrors encountered:\n"
    for i, error in enumerate(errors, 1):
        error_msg += f"{i}. {error}\n"
    
    error_msg += f"\nFile size: {os.path.getsize(file_path)} bytes"
    error_msg += f"\nThis PDF may be corrupted, password-protected, or use unsupported formatting."
    error_msg += f"\nFor PII detection, please convert to a text file or different PDF format."
    
    logger.error("All PDF extraction methods failed")
    return error_msg

def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX files"""
    if not DOCX_AVAILABLE:
        return "[DOCX file - python-docx not installed. Install with: pip install python-docx]"
    
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text += " | ".join(row_text) + "\n"
            text += "--- End Table ---\n\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return f"[DOCX extraction failed: {str(e)}]"

def extract_excel_text(file_path: str) -> str:
    """Extract text from Excel files"""
    if not EXCEL_AVAILABLE:
        return "[Excel file - openpyxl not installed. Install with: pip install openpyxl]"
    
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"=== Sheet: {sheet_name} ===\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None:
                        row_text.append(str(cell))
                if row_text and any(cell.strip() for cell in row_text if isinstance(cell, str)):
                    text += " | ".join(row_text) + "\n"
            text += "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Excel extraction error: {e}")
        return f"[Excel extraction failed: {str(e)}]"

def extract_csv_text(file_path: str) -> str:
    """Extract text from CSV files"""
    try:
        text = ""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            # Try to detect delimiter
            sample = file.read(1024)
            file.seek(0)
            sniffer = csv.Sniffer()
            try:
                delimiter = sniffer.sniff(sample).delimiter
            except:
                delimiter = ','
            
            csv_reader = csv.reader(file, delimiter=delimiter)
            for row_num, row in enumerate(csv_reader):
                text += " | ".join(str(cell) for cell in row) + "\n"
                # Limit rows for very large CSV files
                if row_num > 1000:
                    text += "... (file truncated at 1000 rows for PII analysis)\n"
                    break
        
        return text.strip()
    except Exception as e:
        logger.error(f"CSV extraction error: {e}")
        return f"[CSV extraction failed: {str(e)}]"

def extract_text_file(file_path: str) -> str:
    """Extract text from plain text files"""
    try:
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    # Limit very large text files
                    if len(content) > 100000:  # 100KB text limit
                        content = content[:100000] + "\n... (file truncated for PII analysis)"
                    return content
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try binary mode
        with open(file_path, 'rb') as file:
            content = file.read()
            return content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return f"[Text extraction failed: {str(e)}]"

def extract_json_text(file_path: str) -> str:
    """Extract text from JSON files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        # Convert JSON to readable text
        return json.dumps(data, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"JSON extraction error: {e}")
        return f"[JSON extraction failed: {str(e)}]"

def get_file_extractor(filename: str):
    """Determine which extractor to use based on file extension"""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_pdf_text
    elif filename_lower.endswith('.docx'):
        return extract_docx_text
    elif filename_lower.endswith(('.xlsx', '.xlsm', '.xls')):
        return extract_excel_text
    elif filename_lower.endswith('.csv'):
        return extract_csv_text
    elif filename_lower.endswith('.json'):
        return extract_json_text
    elif filename_lower.endswith(('.txt', '.log', '.md', '.py', '.js', '.css', '.html', '.sql', '.yml', '.yaml', '.xml')):
        return extract_text_file
    else:
        return None

@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    """
    Extract text from uploaded file
    """
    try:
        logger.info(f"Processing file: {file.filename} (type: {file.content_type}, size: {file.size} bytes)")
        
        # Check file size (50MB limit)
        max_size = 50 * 1024 * 1024  # 50MB
        if file.size and file.size > max_size:
            raise HTTPException(
                status_code=413, 
                detail=f"File too large: {file.size} bytes (max: {max_size} bytes)"
            )
        
        # Get the appropriate extractor
        extractor = get_file_extractor(file.filename)
        
        if extractor is None:
            return JSONResponse(
                content={
                    "success": False,
                    "extracted_text": f"[Unsupported file type: {file.filename}] - Manual review recommended for PII",
                    "filename": file.filename,
                    "supported": False,
                    "file_size": file.size
                }
            )
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text using the appropriate extractor
            extracted_text = extractor(temp_file_path)
            
            # Check if extraction actually worked
            if extracted_text.startswith("[") and "failed" in extracted_text.lower():
                logger.warning(f"Extraction failed for {file.filename}: {extracted_text[:100]}...")
                success = False
            else:
                logger.info(f"Successfully extracted {len(extracted_text)} characters from {file.filename}")
                success = True
            
            return JSONResponse(
                content={
                    "success": success,
                    "extracted_text": extracted_text,
                    "filename": file.filename,
                    "supported": True,
                    "file_size": file.size,
                    "text_length": len(extracted_text)
                }
            )
        
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {e}")
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "extracted_text": f"[Error processing file: {file.filename}] - {str(e)}",
                "filename": file.filename,
                "supported": False,
                "error": str(e)
            }
        )

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    dependencies = {
        "PyPDF2": PDF_AVAILABLE,
        "python-docx": DOCX_AVAILABLE, 
        "openpyxl": EXCEL_AVAILABLE,
        "csv": CSV_AVAILABLE,
        "json": JSON_AVAILABLE
    }
    
    return {
        "status": "healthy",
        "message": "File extraction API is running",
        "dependencies": dependencies,
        "available_formats": [
            "PDF (.pdf)" if PDF_AVAILABLE else "PDF (.pdf) - install PyPDF2",
            "Word (.docx)" if DOCX_AVAILABLE else "Word (.docx) - install python-docx", 
            "Excel (.xlsx, .xls)" if EXCEL_AVAILABLE else "Excel (.xlsx, .xls) - install openpyxl",
            "CSV (.csv)",
            "JSON (.json)",
            "Text files (.txt, .py, .js, .css, .html, .sql, etc.)"
        ]
    }

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "name": "File Text Extraction API",
        "version": "1.0.0",
        "description": "Extract text from files for PII detection",
        "endpoints": {
            "extract": "POST /extract-text",
            "health": "GET /health"
        }
    }

if __name__ == "__main__":
    print("🚀 Starting File Text Extraction API...")
    print("📋 This API extracts text from files for PII detection")
    print("🌐 Server will be available at: http://localhost:8080")
    print("📖 Health check: http://localhost:8080/health")
    print("📚 API docs: http://localhost:8080/docs")
    
    uvicorn.run(
        app,
        host="127.0.0.1",
        port=8080,
        log_level="info"
    )
